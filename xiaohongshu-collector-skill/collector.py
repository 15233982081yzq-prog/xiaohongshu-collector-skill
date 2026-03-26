"""
小红书笔记收集器 - 核心收集逻辑
"""

import os
import sys
import json
import time
import datetime
import random
import re
import logging
from pathlib import Path
from typing import Dict, List, Optional, Any, Tuple
import subprocess

logger = logging.getLogger(__name__)

class NoteCollector:
    """笔记收集器"""
    
    def __init__(self, config: Dict, skills_dir: Path):
        self.config = config
        self.skills_dir = skills_dir
        
    def search_feeds(self, keyword: str, limit: int = 10) -> List[Dict]:
        """搜索小红书笔记"""
        try:
            # 添加随机延迟
            delay_min = self.config['advanced']['request_delay_min']
            delay_max = self.config['advanced']['request_delay_max']
            delay = random.uniform(delay_min, delay_max)
            logger.info(f"搜索前等待 {delay:.1f} 秒...")
            time.sleep(delay)
            
            cmd = [
                sys.executable, "scripts/cli.py", 
                "search-feeds", 
                "--keyword", keyword,
                "--sort-by", self.config['basic'].get('sort_by', '综合'),
                "--note-type", self.config['basic'].get('note_type', '不限'),
                "--publish-time", self.config['basic'].get('publish_time', '不限')
            ]
            
            logger.info(f"执行搜索: {keyword}")
            result = subprocess.run(
                cmd,
                cwd=self.skills_dir,
                capture_output=True,
                text=True,
                encoding='utf-8',
                timeout=90
            )
            
            if result.returncode == 0:
                data = json.loads(result.stdout)
                feeds = data.get("feeds", [])
                logger.info(f"搜索到 {len(feeds)} 条笔记")
                return feeds[:limit]
            else:
                logger.error(f"搜索失败: {result.stderr}")
                return []
        except Exception as e:
            logger.error(f"搜索笔记失败: {e}")
            return []
    
    def get_feed_detail(self, feed_id: str, xsec_token: str) -> Optional[Dict]:
        """获取笔记详情"""
        try:
            # 添加随机延迟
            delay_min = self.config['advanced']['request_delay_min']
            delay_max = self.config['advanced']['request_delay_max']
            delay = random.uniform(delay_min + 1, delay_max + 2)
            logger.info(f"获取详情前等待 {delay:.1f} 秒...")
            time.sleep(delay)
            
            cmd = [
                sys.executable, "scripts/cli.py",
                "get-feed-detail",
                "--feed-id", feed_id,
                "--xsec-token", xsec_token
            ]
            
            logger.info(f"获取笔记详情: {feed_id}")
            result = subprocess.run(
                cmd,
                cwd=self.skills_dir,
                capture_output=True,
                text=True,
                encoding='utf-8',
                timeout=150
            )
            
            if result.returncode == 0:
                detail_data = json.loads(result.stdout)
                note_data = detail_data.get('note', {})
                
                # 构建标准化的feed数据
                feed_detail = {
                    'id': note_data.get('noteId', feed_id),
                    'title': note_data.get('title', ''),
                    'desc': note_data.get('desc', ''),
                    'content': note_data.get('desc', ''),
                    'user': note_data.get('user', {}),
                    'time': note_data.get('time'),
                    'ipLocation': note_data.get('ipLocation', ''),
                    'type': note_data.get('type', ''),
                    'interactInfo': note_data.get('interactInfo', {}),
                    'url': f"https://www.xiaohongshu.com/explore/{feed_id}"
                }
                
                logger.info(f"获取到笔记详情: {feed_detail.get('title', '无标题')}")
                return feed_detail
            else:
                logger.error(f"获取详情失败: {result.stderr}")
                return None
        except Exception as e:
            logger.error(f"获取笔记详情失败: {e}")
            return None
    
    def filter_content(self, content: str, title: str = "") -> Dict[str, Any]:
        """过滤内容"""
        if not content:
            return {"relevant": False, "score": 0, "reasons": ["内容为空"], "matched_keywords": []}
        
        # 合并标题和内容
        full_text = (title + " " + content).lower()
        
        # 检查排除关键词
        exclude_keywords = self.config['filter']['exclude_keywords']
        for keyword in exclude_keywords:
            if keyword.lower() in full_text:
                return {
                    "relevant": False, 
                    "score": 0, 
                    "reasons": [f"包含排除关键词: {keyword}"],
                    "matched_keywords": []
                }
        
        # 检查包含关键词
        include_keywords = self.config['filter']['include_keywords']
        matched_keywords = []
        match_count = 0
        
        for keyword in include_keywords:
            if keyword.lower() in full_text:
                match_count += 1
                matched_keywords.append(keyword)
        
        # 计算分数
        score = 0
        reasons = []
        
        # 关键词匹配分数
        if match_count >= 3:
            score += 3
            reasons.append(f"匹配多个关键词: {match_count}个")
        elif match_count >= 2:
            score += 2
            reasons.append(f"匹配关键词: {match_count}个")
        elif match_count >= 1:
            score += 1
            reasons.append(f"匹配关键词: {match_count}个")
        
        # 严格检查内容长度
        content_length = len(content)
        min_length = self.config['filter']['min_content_length']
        
        # 如果启用了严格模式，不满足最小字数的直接排除
        strict_mode = self.config['filter'].get('strict_mode', True)
        
        if content_length < min_length:
            if strict_mode:
                # 严格模式：不满足字数直接返回不合格
                return {
                    "relevant": False,
                    "score": 0,
                    "reasons": [f"内容过短: {content_length}字 < 要求{min_length}字"],
                    "matched_keywords": matched_keywords
                }
            else:
                # 非严格模式：只是不加分
                reasons.append(f"内容较短: {content_length}字 < 要求{min_length}字")
        else:
            # 满足字数要求，根据长度加分
            if content_length > min_length * 2:
                score += 10
                reasons.append(f"内容详细: {content_length}字")
            elif content_length > min_length:
                score += 9
                reasons.append(f"内容适中: {content_length}字")
        
        # 格式分数
        if re.search(r'[0-9]+[\.、]\s*[^。]+[？?]', content):
            score += 2
            reasons.append("包含编号的问题列表")
        
        if '回答' in content or '解答' in content or '答案' in content:
            score += 1
            reasons.append("包含回答/解答")
        
        if '经验' in content or '建议' in content or '技巧' in content:
            score += 1
            reasons.append("包含经验分享")
        
        # 判断是否相关
        threshold = self.config['filter']['quality_threshold']
        strict_mode = self.config['filter']['strict_mode']
        
        if strict_mode:
            relevant = score >= threshold and match_count >= 2
        else:
            relevant = score >= threshold or match_count >= 2
        
        return {
            "relevant": relevant,
            "score": score,
            "reasons": reasons,
            "matched_keywords": matched_keywords[:5]
        }
    
    def collect(self, target_count: int, keyword: str) -> List[Dict]:
        """收集笔记"""
        logger.info(f"开始收集，目标: {target_count} 篇笔记，关键词: {keyword}")
        
        collected_feeds = []
        search_attempts = 0
        max_attempts = self.config['advanced']['max_retries']
        
        while len(collected_feeds) < target_count and search_attempts < max_attempts:
            search_attempts += 1
            logger.info(f"第 {search_attempts} 次搜索尝试...")
            
            # 搜索笔记
            feeds = self.search_feeds(keyword, limit=15)
            if not feeds:
                logger.warning("未搜索到笔记")
                continue
            
            # 处理每条笔记
            for i, feed in enumerate(feeds):
                if len(collected_feeds) >= target_count:
                    break
                
                feed_id = feed.get('id')
                xsec_token = feed.get('xsecToken')
                title = feed.get('displayTitle', '')
                
                logger.info(f"处理笔记 {i+1}/{len(feeds)}: {title}")
                
                if not feed_id or not xsec_token:
                    logger.warning("笔记缺少ID或token，跳过")
                    continue
                
                # 获取详情
                detail = self.get_feed_detail(feed_id, xsec_token)
                if not detail:
                    logger.warning("获取详情失败，跳过")
                    continue
                
                # 内容筛选
                content = detail.get('desc', '') or detail.get('content', '')
                filter_result = self.filter_content(content, detail.get('title', ''))
                
                if filter_result['relevant']:
                    # 添加筛选信息
                    detail['filter_info'] = filter_result
                    detail['collected_time'] = datetime.datetime.now().isoformat()
                    
                    collected_feeds.append(detail)
                    logger.info(f"✅ 收集到第 {len(collected_feeds)} 篇笔记")
                    logger.info(f"   评分: {filter_result['score']}, 原因: {', '.join(filter_result['reasons'])}")
                else:
                    logger.info(f"❌ 筛选不通过: {filter_result['reasons']}")
            
            # 如果还不够，等待后继续
            if len(collected_feeds) < target_count:
                wait_time = random.uniform(10.0, 20.0)
                logger.info(f"已收集 {len(collected_feeds)}/{target_count} 篇，等待 {wait_time:.1f} 秒...")
                time.sleep(wait_time)
        
        # 按评分排序
        collected_feeds.sort(key=lambda x: x.get('filter_info', {}).get('score', 0), reverse=True)
        
        logger.info(f"收集完成，共收集到 {len(collected_feeds)} 篇笔记")
        return collected_feeds[:target_count]
    
    def check_login(self) -> bool:
        """检查小红书登录状态"""
        try:
            cmd = [sys.executable, "scripts/cli.py", "check-login"]
            result = subprocess.run(
                cmd,
                cwd=self.skills_dir,
                capture_output=True,
                text=True,
                encoding='utf-8',
                timeout=30
            )
            
            if result.returncode == 0:
                data = json.loads(result.stdout)
                is_logged_in = data.get('logged_in', False)
                logger.info(f"登录状态检查: {'已登录' if is_logged_in else '未登录'}")
                return is_logged_in
            else:
                logger.warning(f"登录检查失败: {result.stderr}")
                return False
                
        except Exception as e:
            logger.error(f"检查登录状态失败: {e}")
            return False
    
    def save_to_word(self, feeds: List[Dict]) -> bool:
        """保存到Word文档"""
        try:
            import docx
            from docx.shared import Inches
            
            # 获取输出路径
            output_dir = Path(self.config['basic']['output_dir'])
            output_dir.mkdir(parents=True, exist_ok=True)
            
            filename = self.config['basic']['output_filename']
            
            # 替换模板变量
            if '{date}' in filename:
                date_str = datetime.datetime.now().strftime('%Y%m%d')
                filename = filename.replace('{date}', date_str)
            
            if '{keyword}' in filename:
                keyword = self.config['basic']['search_keyword']
                filename = filename.replace('{keyword}', keyword)
            
            if '{count}' in filename:
                count = self.config['basic']['collect_count']
                filename = filename.replace('{count}', str(count))
            
            # 添加时间戳确保文件名唯一
            name_parts = filename.split('.')
            if len(name_parts) > 1:
                timestamp = datetime.datetime.now().strftime('%Y%m%d_%H%M%S')
                name_parts[-2] = f"{name_parts[-2]}_{timestamp}"
                filename = '.'.join(name_parts)
            
            word_file = output_dir / filename
            
            # 创建文档
            doc = docx.Document()
            doc.add_heading('小红书笔记收集结果', 0)
            
            # 添加收集信息
            doc.add_paragraph(f"收集时间: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            doc.add_paragraph(f"收集数量: {len(feeds)} 篇")
            doc.add_paragraph(f"搜索关键词: {self.config['basic']['search_keyword']}")
            doc.add_paragraph()
            
            # 添加每篇笔记
            for i, feed in enumerate(feeds, 1):
                doc.add_heading(f"{i}. {feed.get('title', '无标题')}", 2)
                
                # 添加基本信息
                if feed.get('time'):
                    publish_time = datetime.datetime.fromtimestamp(feed['time']/1000).strftime('%Y-%m-%d %H:%M:%S')
                    doc.add_paragraph(f"发布时间: {publish_time}")
                
                if feed.get('user', {}).get('nickname'):
                    doc.add_paragraph(f"作者: {feed['user']['nickname']}")
                
                if feed.get('ipLocation'):
                    doc.add_paragraph(f"位置: {feed['ipLocation']}")
                
                # 添加内容
                content = feed.get('desc') or feed.get('content', '')
                if content:
                    doc.add_paragraph("内容:")
                    paragraphs = content.split('\n')
                    for para in paragraphs:
                        if para.strip():
                            doc.add_paragraph(para.strip())
                
                # 添加筛选信息
                if feed.get('filter_info'):
                    filter_info = feed['filter_info']
                    doc.add_paragraph(f"筛选评分: {filter_info.get('score', 0)}")
                    doc.add_paragraph(f"入选原因: {', '.join(filter_info.get('reasons', []))}")
                
                # 添加分隔线
                if i < len(feeds):
                    doc.add_paragraph("-" * 50)
                    doc.add_paragraph()
            
            # 保存文档
            doc.save(str(word_file))
            logger.info(f"已保存到Word文档: {word_file}")
            return True
            
        except Exception as e:
            logger.error(f"保存到Word文档失败: {e}")
            return False
    
    def save_to_json(self, feeds: List[Dict]) -> bool:
        """保存到JSON文件"""
        try:
            import json
            
            # 获取输出路径
            output_dir = Path(self.config['basic']['output_dir'])
            output_dir.mkdir(parents=True, exist_ok=True)
            
            # 构建文件名
            keyword = self.config['basic']['search_keyword']
            timestamp = datetime.datetime.now().strftime('%Y%m%d_%H%M%S')
            json_file = output_dir / f"小红书笔记_{keyword}_{timestamp}.json"
            
            # 准备数据
            data = {
                "metadata": {
                    "collect_time": datetime.datetime.now().isoformat(),
                    "search_keyword": keyword,
                    "collected_count": len(feeds),
                    "config": self.config
                },
                "notes": feeds
            }
            
            # 保存JSON文件
            with open(json_file, 'w', encoding='utf-8') as f:
                json.dump(data, f, indent=2, ensure_ascii=False)
            
            logger.info(f"已保存到JSON文件: {json_file}")
            return True
            
        except Exception as e:
            logger.error(f"保存到JSON文件失败: {e}")
            return False
    
    def collect_and_save(self) -> bool:
        """收集并保存笔记"""
        try:
            # 获取配置参数
            keyword = self.config['basic']['search_keyword']
            count = self.config['basic']['collect_count']
            
            # 收集笔记
            feeds = self.collect(count, keyword)
            
            if not feeds:
                logger.warning("未收集到任何笔记")
                return False
            
            # 保存到Word
            word_success = self.save_to_word(feeds)
            
            # 保存到JSON
            json_success = self.save_to_json(feeds)
            
            if word_success or json_success:
                logger.info(f"成功收集并保存 {len(feeds)} 篇笔记")
                return True
            else:
                logger.error("保存笔记失败")
                return False
                
        except Exception as e:
            logger.error(f"收集和保存失败: {e}")
            return False